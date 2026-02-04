#!/usr/bin/env python3
"""
creative_gemini_visual_annotator.py

Creates "circled / arrowed" annotated images that highlight *visible* creative elements
likely associated with higher/lower performance (CTR / CVR / ATC), plus:
  - a Markdown + JSON callout file with longer explanations for each numbered annotation
  - a well-formatted Word doc (.docx) embedding the annotated images and explanations

This script is designed to run AFTER (or alongside) your metric + Gemini creative analysis flow.
It can optionally consume:
  - --analysis-jsonl  (from creative_gemini_creative_report.py, per_asset_analysis.jsonl)
  - --signals-csv     (from creative_gemini_creative_report.py, feature_signals.csv)

But it can also run without them.

Key idea:
- Gemini does NOT directly "draw" circles/arrows reliably in the API.
- Instead, Gemini returns structured callouts (with normalized bounding boxes),
  and *this script* draws circles/arrows/labels onto the images.

For videos:
- We extract a representative frame (midpoint) and annotate that frame.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import mimetypes
import os
import subprocess
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from pydantic import BaseModel, Field, ValidationError

from google import genai
from google.genai import types

# ----------------------------
# Defaults
# ----------------------------

DEFAULT_MODEL_ANNOTATE = os.getenv("GEMINI_MODEL_ANNOTATE", "gemini-3-flash-preview")
FALLBACK_MODEL_ANNOTATE = "gemini-2.5-flash"

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp"}
VIDEO_EXTS = {".mp4", ".mov", ".m4v", ".webm", ".mkv"}

DEFAULT_METRICS = ["click_through_rate", "conversion_rate", "add_to_cart_rate"]


# ----------------------------
# Schema for Gemini structured output
# ----------------------------

class BBox(BaseModel):
    # normalized coords in [0,1], left/top/right/bottom
    x1: float = Field(ge=0.0, le=1.0)
    y1: float = Field(ge=0.0, le=1.0)
    x2: float = Field(ge=0.0, le=1.0)
    y2: float = Field(ge=0.0, le=1.0)

class Callout(BaseModel):
    id: str = Field(description="Unique id like A1, A2, B1, etc.")
    target: str = Field(description="Which image the box refers to: 'A' or 'B'.")
    metric_focus: List[str] = Field(description="One or more of: click_through_rate, conversion_rate, add_to_cart_rate")
    direction: str = Field(description="Either 'increases' or 'decreases' relative to the metric_focus.")
    label: str = Field(description="Short label for the callout (3-7 words).")
    why: str = Field(description="1-3 sentences explaining the likely impact; must reference the metric_focus.")
    box: BBox = Field(description="Normalized bounding box for the highlighted region.")
    shape: str = Field(default="circle", description="circle or arrow (circle recommended).")

class PairwiseAnnotation(BaseModel):
    metric: str
    a_filename: str
    b_filename: str
    callouts: List[Callout] = Field(description="Callouts across A and B")
    notes: Optional[str] = None

PAIR_SCHEMA = PairwiseAnnotation.model_json_schema()


# ----------------------------
# Helpers
# ----------------------------

def eprint(*args: Any) -> None:
    print(*args, file=sys.stderr)

def safe_mkdir(p: Path) -> None:
    p.mkdir(parents=True, exist_ok=True)

def norm_filename(name: str) -> str:
    return Path(name).name.strip().lower()

def fingerprint_file(path: Path, head_bytes: int = 1024 * 1024) -> str:
    st = path.stat()
    h = hashlib.sha1()
    h.update(str(st.st_size).encode("utf-8"))
    h.update(str(int(st.st_mtime)).encode("utf-8"))
    with path.open("rb") as f:
        h.update(f.read(head_bytes))
    return h.hexdigest()

def guess_mime(path: Path) -> str:
    mime, _ = mimetypes.guess_type(str(path))
    if mime:
        return mime
    if path.suffix.lower() in IMAGE_EXTS:
        return "image/jpeg"
    if path.suffix.lower() in VIDEO_EXTS:
        return "video/mp4"
    return "application/octet-stream"

def retry_call(fn, *, max_tries: int = 4, base_sleep: float = 1.2):
    last_err = None
    for i in range(max_tries):
        try:
            return fn()
        except Exception as e:
            last_err = e
            sleep_s = base_sleep * (2 ** i)
            eprint(f"[warn] API call failed (try {i+1}/{max_tries}): {e}. Sleeping {sleep_s:.1f}s")
            time.sleep(sleep_s)
    raise last_err

def make_client(args: argparse.Namespace) -> genai.Client:
    if args.vertexai:
        return genai.Client(vertexai=True, project=args.project, location=args.location)
    if args.api_key:
        return genai.Client(api_key=args.api_key)
    return genai.Client()


def list_media_files(media_dir: Path, recursive: bool) -> List[Path]:
    patterns = ["*"] if not recursive else ["**/*"]
    out: List[Path] = []
    for pat in patterns:
        for p in media_dir.glob(pat):
            if p.is_file():
                ext = p.suffix.lower()
                if ext in IMAGE_EXTS or ext in VIDEO_EXTS:
                    out.append(p)
    out.sort()
    return out


def coerce_numeric(df: pd.DataFrame, cols: List[str]) -> pd.DataFrame:
    for c in cols:
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")
    return df


def ensure_rates(df: pd.DataFrame) -> pd.DataFrame:
    # CTR
    if "click_through_rate" not in df.columns:
        df["click_through_rate"] = np.where(df["impressions"] > 0, df["clicks"] / df["impressions"], np.nan)
    # CVR
    if "conversion_rate" not in df.columns:
        df["conversion_rate"] = np.where(df["clicks"] > 0, df["conversions"] / df["clicks"], np.nan)
    # ATC rate
    if "add_to_cart_rate" not in df.columns:
        df["add_to_cart_rate"] = np.where(df["clicks"] > 0, df["add_to_carts"] / df["clicks"], np.nan)
    return df


def load_image_bytes(path: Path, max_side: int = 1400, jpeg_quality: int = 92) -> bytes:
    img = Image.open(path).convert("RGB")
    w, h = img.size
    scale = min(1.0, max_side / float(max(w, h)))
    if scale < 1.0:
        img = img.resize((int(w * scale), int(h * scale)))
    import io
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=jpeg_quality, optimize=True)
    return buf.getvalue()


def get_video_duration_seconds(path: Path) -> Optional[float]:
    try:
        cmd = [
            "ffprobe", "-v", "error",
            "-show_entries", "format=duration",
            "-of", "default=noprint_wrappers=1:nokey=1",
            str(path),
        ]
        res = subprocess.run(cmd, capture_output=True, text=True, check=False)
        if res.returncode == 0 and res.stdout.strip():
            return float(res.stdout.strip())
    except Exception:
        pass
    return None


def extract_mid_frame(video_path: Path, out_frame: Path) -> Optional[Path]:
    """
    Extract a midpoint frame using ffmpeg if available; fall back to OpenCV.
    """
    safe_mkdir(out_frame.parent)

    dur = get_video_duration_seconds(video_path)
    t = None
    if dur and dur > 0:
        t = max(0.0, dur * 0.5)

    # ffmpeg approach
    try:
        if t is None:
            # take first frame
            cmd = ["ffmpeg", "-y", "-i", str(video_path), "-frames:v", "1", str(out_frame)]
        else:
            cmd = ["ffmpeg", "-y", "-ss", f"{t:.3f}", "-i", str(video_path), "-frames:v", "1", str(out_frame)]
        res = subprocess.run(cmd, capture_output=True, text=True, check=False)
        if res.returncode == 0 and out_frame.exists():
            return out_frame
    except FileNotFoundError:
        pass
    except Exception:
        pass

    # OpenCV fallback
    try:
        import cv2
        cap = cv2.VideoCapture(str(video_path))
        if not cap.isOpened():
            return None
        frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
        idx = frame_count // 2 if frame_count > 0 else 0
        cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
        ok, frame = cap.read()
        cap.release()
        if not ok:
            return None
        # BGR -> RGB
        frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        img = Image.fromarray(frame)
        img.save(out_frame)
        return out_frame
    except Exception:
        return None


def get_representative_image(media_path: Path, frames_dir: Path) -> Tuple[Path, str]:
    """
    Returns (image_path, kind) where kind is 'image' or 'video_frame'.
    """
    ext = media_path.suffix.lower()
    if ext in IMAGE_EXTS:
        return media_path, "image"
    if ext in VIDEO_EXTS:
        frame_path = frames_dir / f"{media_path.stem}__midframe.jpg"
        if frame_path.exists():
            return frame_path, "video_frame"
        got = extract_mid_frame(media_path, frame_path)
        if got is None:
            raise RuntimeError(f"Could not extract frame from video: {media_path}")
        return got, "video_frame"
    raise ValueError(f"Unsupported media type: {media_path}")


def load_analysis_jsonl(path: Path) -> Dict[str, Dict[str, Any]]:
    """
    Load per_asset_analysis.jsonl into dict keyed by normalized filename.
    """
    out: Dict[str, Dict[str, Any]] = {}
    if not path.exists():
        return out
    with path.open("r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
                fn = obj.get("filename") or obj.get("file_name") or obj.get("name")
                if fn:
                    out[norm_filename(fn)] = obj
            except Exception:
                continue
    return out


def load_signals(signals_csv: Optional[Path]) -> Dict[str, List[Dict[str, Any]]]:
    """
    Load feature_signals.csv and return top signals per metric.
    """
    out: Dict[str, List[Dict[str, Any]]] = {}
    if signals_csv is None or (not signals_csv.exists()):
        return out
    try:
        df = pd.read_csv(signals_csv)
        if df.empty or "metric" not in df.columns:
            return out
        for metric in df["metric"].dropna().unique().tolist():
            sub = df[df["metric"] == metric].copy()
            # prefer strongest signals
            if "value" in sub.columns:
                sub["abs_value"] = sub["value"].abs()
                sub = sub.sort_values("abs_value", ascending=False).drop(columns=["abs_value"], errors="ignore")
            out[str(metric)] = sub.head(10).to_dict(orient="records")
    except Exception:
        return out
    return out


def format_metric(v: Any) -> str:
    try:
        if v is None or (isinstance(v, float) and np.isnan(v)):
            return "n/a"
        return f"{float(v):.4f}"
    except Exception:
        return str(v)


# ----------------------------
# Drawing annotations
# ----------------------------

def _load_font(size: int = 26) -> ImageFont.FreeTypeFont:
    # Try a reasonable default font; fall back to PIL default.
    for p in [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
    ]:
        try:
            if Path(p).exists():
                return ImageFont.truetype(p, size=size)
        except Exception:
            pass
    return ImageFont.load_default()

def draw_callouts(
    image_path: Path,
    callouts: List[Callout],
    *,
    target: str,
    out_path: Path,
    circle_color: Tuple[int, int, int],
    text_color: Tuple[int, int, int] = (255, 255, 255),
) -> None:
    img = Image.open(image_path).convert("RGB")
    W, H = img.size
    draw = ImageDraw.Draw(img)
    font = _load_font(26)

    for c in callouts:
        if c.target != target:
            continue
        x1 = int(c.box.x1 * W)
        y1 = int(c.box.y1 * H)
        x2 = int(c.box.x2 * W)
        y2 = int(c.box.y2 * H)

        # Expand a bit for nicer circle
        pad = int(0.01 * max(W, H))
        x1p, y1p, x2p, y2p = max(0, x1 - pad), max(0, y1 - pad), min(W - 1, x2 + pad), min(H - 1, y2 + pad)

        # circle/ellipse around region
        draw.ellipse([x1p, y1p, x2p, y2p], outline=circle_color, width=max(4, int(0.004 * max(W, H))))

        # label background box
        label = c.id
        tx, ty = x1p, max(0, y1p - int(0.04 * H))
        # measure text
        bbox = draw.textbbox((tx, ty), label, font=font)
        bx1, by1, bx2, by2 = bbox
        pad2 = 6
        draw.rectangle([bx1 - pad2, by1 - pad2, bx2 + pad2, by2 + pad2], fill=circle_color)
        draw.text((tx, ty), label, font=font, fill=text_color)

    safe_mkdir(out_path.parent)
    img.save(out_path)


# ----------------------------
# Gemini annotation call
# ----------------------------

def annotate_pairwise(
    client: genai.Client,
    model: str,
    metric: str,
    a_img_path: Path,
    b_img_path: Path,
    a_ctx: Dict[str, Any],
    b_ctx: Dict[str, Any],
    top_signals: List[Dict[str, Any]],
) -> PairwiseAnnotation:
    """
    Send two images + context to Gemini; get structured callouts with boxes.
    """
    a_bytes = load_image_bytes(a_img_path)
    b_bytes = load_image_bytes(b_img_path)

    # Keep context short and deterministic
    def compact_ctx(ctx: Dict[str, Any]) -> Dict[str, Any]:
        keep = [
            "asset_type", "brand_name", "product_name", "product_category",
            "messaging_summary", "offer_or_promo", "tone",
            "cta_present", "cta_text",
            "background_style", "background_description",
            "camera_motion_type", "zoom_behavior",
            "video_length_seconds",
            "message_clarity_0_5", "text_readability_0_5", "cta_strength_0_5",
            "brand_readability_0_5", "product_prominence_0_5", "lighting_quality_0_5",
            "background_distraction_0_5", "people_present",
            "tags",
        ]
        out = {}
        for k in keep:
            if k in ctx:
                out[k] = ctx[k]
        return out

    # Instruction prompt
    prompt = (
        "You are a performance creative analyst.\n"
        "I will provide two creatives (A and B) and their performance metrics.\n"
        "Your job is to point to *visible* differences that plausibly explain why one performs better on the chosen metric.\n\n"
        f"TARGET METRIC: {metric}\n\n"
        "Return ONLY a JSON object that matches the schema.\n"
        "Rules:\n"
        "- Provide 4-8 callouts total (across both images).\n"
        "- Each callout MUST include a normalized bounding box [0..1].\n"
        "- The highlighted region must be something visually locatable: CTA text, offer badge, brand name, product, background element, contrast/color block, etc.\n"
        "- Be careful: do not invent text that isn't visible.\n"
        "- Do not claim causality; use language like 'associated with' / 'could contribute'.\n"
        "- If the chosen metric is click_through_rate, prioritize attention/scroll-stopping elements.\n"
        "- If conversion_rate or add_to_cart_rate, prioritize clarity, trust, offer, product understanding, and friction reduction.\n\n"
        "Performance context (A is higher on this metric, B is lower):\n"
        f"A metrics: {json.dumps(a_ctx.get('metrics', {}), ensure_ascii=False)}\n"
        f"B metrics: {json.dumps(b_ctx.get('metrics', {}), ensure_ascii=False)}\n\n"
        "Dataset-level signals (top associations for this metric):\n"
        f"{json.dumps(top_signals, ensure_ascii=False)}\n\n"
        "Creative analysis context (if available):\n"
        f"A creative context: {json.dumps(compact_ctx(a_ctx.get('analysis', {})), ensure_ascii=False)}\n"
        f"B creative context: {json.dumps(compact_ctx(b_ctx.get('analysis', {})), ensure_ascii=False)}\n\n"
        "Now generate the structured callouts with bounding boxes for BOTH images.\n"
    )

    contents = types.Content(parts=[
        types.Part(inline_data=types.Blob(data=a_bytes, mime_type="image/jpeg")),
        types.Part(inline_data=types.Blob(data=b_bytes, mime_type="image/jpeg")),
        types.Part(text=prompt),
    ])

    config = {
        "response_mime_type": "application/json",
        "response_json_schema": PAIR_SCHEMA,
    }

    def _call(m: str):
        return client.models.generate_content(model=m, contents=contents, config=config)

    try:
        resp = retry_call(lambda: _call(model))
    except Exception as e:
        eprint(f"[warn] annotate failed with model={model}: {e}. Falling back to {FALLBACK_MODEL_ANNOTATE}")
        resp = retry_call(lambda: _call(FALLBACK_MODEL_ANNOTATE))

    text = resp.text or "{}"
    try:
        obj = PairwiseAnnotation.model_validate_json(text)
    except ValidationError:
        # fallback: parse json, then validate dict
        obj = PairwiseAnnotation.model_validate(json.loads(text))

    return obj


# ----------------------------
# Word doc generation
# ----------------------------

def build_docx(
    out_docx: Path,
    title: str,
    dataset_summary: Dict[str, Any],
    metric_sections: List[Dict[str, Any]],
) -> None:
    """
    Create a well-formatted Word doc with annotated images and callouts.
    """
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # Basic styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(11)

    # Title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Summary
    doc.add_heading("Dataset summary", level=1)
    for k, v in dataset_summary.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_paragraph()

    for sec in metric_sections:
        metric = sec["metric"]
        doc.add_heading(f"Metric: {metric}", level=1)

        # Add a small table for A/B
        a = sec["A"]
        b = sec["B"]
        tbl = doc.add_table(rows=2, cols=2)
        tbl.style = "Table Grid"

        tbl.cell(0, 0).text = f"A (Higher {metric})\n{a['file_name']}\nCTR={a['ctr']}, CVR={a['cvr']}, ATC={a['atc']}"
        tbl.cell(0, 1).text = f"B (Lower {metric})\n{b['file_name']}\nCTR={b['ctr']}, CVR={b['cvr']}, ATC={b['atc']}"

        # images row
        tbl.cell(1, 0).paragraphs[0].add_run().add_picture(str(a["annotated_image"]), width=Inches(3.2))
        tbl.cell(1, 1).paragraphs[0].add_run().add_picture(str(b["annotated_image"]), width=Inches(3.2))

        doc.add_paragraph()

        doc.add_heading("Callouts", level=2)

        # Group callouts
        callouts: List[Dict[str, Any]] = sec["callouts"]
        a_callouts = [c for c in callouts if c["target"] == "A"]
        b_callouts = [c for c in callouts if c["target"] == "B"]

        doc.add_paragraph("A (higher):", style=None).runs[0].bold = True
        for c in a_callouts:
            doc.add_paragraph(f"{c['id']} — {c['label']}: {c['why']}", style="List Bullet")

        doc.add_paragraph("B (lower):", style=None).runs[0].bold = True
        for c in b_callouts:
            doc.add_paragraph(f"{c['id']} — {c['label']}: {c['why']}", style="List Bullet")

        if sec.get("notes"):
            doc.add_paragraph()
            doc.add_paragraph(f"Notes: {sec['notes']}")

        doc.add_page_break()

    safe_mkdir(out_docx.parent)
    doc.save(str(out_docx))


# ----------------------------
# Main
# ----------------------------

def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--media-dir", required=True, help="Folder containing images/videos")
    ap.add_argument("--metrics-csv", required=True, help="CSV keyed by filename with metrics")
    ap.add_argument("--out-dir", required=True, help="Output folder")
    ap.add_argument("--recursive", action="store_true", help="Recursively scan media-dir")
    ap.add_argument("--file-col", default="file_name", help="CSV column matching filenames (default: file_name)")

    ap.add_argument("--metrics", nargs="*", default=DEFAULT_METRICS, help="Metrics to annotate (default: CTR/CVR/ATC)")
    ap.add_argument("--top-k", type=int, default=1, help="How many top and bottom creatives to consider (default: 1). Pairwise uses top-1 vs bottom-1.")
    ap.add_argument("--pairwise", action="store_true", help="Generate pairwise top-vs-bottom annotations per metric (recommended).")

    ap.add_argument("--analysis-jsonl", default=None, help="Optional per_asset_analysis.jsonl from the other script")
    ap.add_argument("--signals-csv", default=None, help="Optional feature_signals.csv from the other script")

    ap.add_argument("--model-annotate", default=DEFAULT_MODEL_ANNOTATE, help="Gemini model for annotation callouts")
    ap.add_argument("--image-max-side", type=int, default=1400, help="Downscale images for Gemini upload")
    ap.add_argument("--no-docx", action="store_true", help="Skip Word doc generation")
    ap.add_argument("--docx-name", default="annotated_creative_report.docx", help="Output docx filename")

    # Auth
    ap.add_argument("--api-key", default=None, help="Explicit API key (else use GEMINI_API_KEY/GOOGLE_API_KEY env vars)")
    ap.add_argument("--vertexai", action="store_true", help="Use Vertex AI backend")
    ap.add_argument("--project", default=os.getenv("GOOGLE_CLOUD_PROJECT"), help="Vertex project (if --vertexai)")
    ap.add_argument("--location", default=os.getenv("GOOGLE_CLOUD_LOCATION", "us-central1"), help="Vertex location (if --vertexai)")

    args = ap.parse_args()

    out_dir = Path(args.out_dir)
    safe_mkdir(out_dir)
    cache_dir = out_dir / "cache"
    frames_dir = out_dir / "frames"
    ann_dir = out_dir / "annotations"
    img_dir = out_dir / "annotated_images"
    safe_mkdir(cache_dir)
    safe_mkdir(frames_dir)
    safe_mkdir(ann_dir)
    safe_mkdir(img_dir)

    client = make_client(args)

    # Load metrics
    dfm = pd.read_csv(args.metrics_csv)
    if args.file_col not in dfm.columns:
        eprint(f"[error] CSV missing file column '{args.file_col}'. Found: {list(dfm.columns)}")
        return 2

    # Normalize + coerce
    dfm = dfm.rename(columns={args.file_col: "file_name"})
    dfm["file_name_norm"] = dfm["file_name"].astype(str).apply(norm_filename)

    base_cols = ["impressions", "clicks", "cost", "conversions", "add_to_carts",
                 "click_through_rate", "conversion_rate", "add_to_cart_rate"]
    dfm = coerce_numeric(dfm, [c for c in base_cols if c in dfm.columns])

    # Ensure base columns exist
    for c in ["impressions", "clicks", "cost", "conversions", "add_to_carts"]:
        if c not in dfm.columns:
            dfm[c] = np.nan

    dfm = ensure_rates(dfm)

    # Discover media and join
    media_files = list_media_files(Path(args.media_dir), args.recursive)
    if not media_files:
        eprint("[error] No media files found.")
        return 2

    rows: List[Dict[str, Any]] = []
    missing = 0
    for p in media_files:
        key = norm_filename(p.name)
        match = dfm[dfm["file_name_norm"] == key]
        if match.empty:
            missing += 1
            continue
        r = match.iloc[0].to_dict()
        r["media_path"] = str(p)
        rows.append(r)
    if not rows:
        eprint("[error] No media matched the CSV filenames.")
        return 2
    if missing:
        eprint(f"[warn] {missing} media files had no matching CSV row (skipped).")

    df = pd.DataFrame(rows)

    # Load optional analysis + signals
    analysis_map: Dict[str, Dict[str, Any]] = {}
    if args.analysis_jsonl:
        analysis_map = load_analysis_jsonl(Path(args.analysis_jsonl))

    signals_map = load_signals(Path(args.signals_csv) if args.signals_csv else None)

    # Dataset summary
    dataset_summary = {
        "n_creatives": int(len(df)),
        "metrics_csv": str(Path(args.metrics_csv).resolve()),
        "media_dir": str(Path(args.media_dir).resolve()),
        "metrics_used": ", ".join(args.metrics),
        "ctr_mean": format_metric(df["click_through_rate"].mean()),
        "cvr_mean": format_metric(df["conversion_rate"].mean()),
        "atc_mean": format_metric(df["add_to_cart_rate"].mean()),
    }

    metric_sections: List[Dict[str, Any]] = []

    # Pairwise: top-1 vs bottom-1 per metric
    for metric in args.metrics:
        if metric not in df.columns:
            eprint(f"[warn] metric '{metric}' not found; skipping.")
            continue

        dfx = df.dropna(subset=[metric]).copy()
        if dfx.empty or len(dfx) < 2:
            eprint(f"[warn] not enough rows with metric '{metric}' to compare.")
            continue

        dfx = dfx.sort_values(metric, ascending=False)
        top = dfx.iloc[0]
        bot = dfx.iloc[-1]

        # Prepare representative images
        top_media = Path(top["media_path"])
        bot_media = Path(bot["media_path"])
        top_img, top_kind = get_representative_image(top_media, frames_dir)
        bot_img, bot_kind = get_representative_image(bot_media, frames_dir)

        # Context bundles
        top_key = norm_filename(top_media.name)
        bot_key = norm_filename(bot_media.name)

        a_ctx = {
            "metrics": {
                "click_through_rate": float(top.get("click_through_rate", np.nan)),
                "conversion_rate": float(top.get("conversion_rate", np.nan)),
                "add_to_cart_rate": float(top.get("add_to_cart_rate", np.nan)),
                "impressions": float(top.get("impressions", np.nan)),
                "clicks": float(top.get("clicks", np.nan)),
                "conversions": float(top.get("conversions", np.nan)),
                "add_to_carts": float(top.get("add_to_carts", np.nan)),
                "cost": float(top.get("cost", np.nan)),
                "role": f"higher_{metric}",
                "media_kind": top_kind,
            },
            "analysis": analysis_map.get(top_key, {}),
        }
        b_ctx = {
            "metrics": {
                "click_through_rate": float(bot.get("click_through_rate", np.nan)),
                "conversion_rate": float(bot.get("conversion_rate", np.nan)),
                "add_to_cart_rate": float(bot.get("add_to_cart_rate", np.nan)),
                "impressions": float(bot.get("impressions", np.nan)),
                "clicks": float(bot.get("clicks", np.nan)),
                "conversions": float(bot.get("conversions", np.nan)),
                "add_to_carts": float(bot.get("add_to_carts", np.nan)),
                "cost": float(bot.get("cost", np.nan)),
                "role": f"lower_{metric}",
                "media_kind": bot_kind,
            },
            "analysis": analysis_map.get(bot_key, {}),
        }

        # Cache key
        fp_a = fingerprint_file(top_img)
        fp_b = fingerprint_file(bot_img)
        cache_path = cache_dir / f"pair_{metric}__{top_media.stem}__{bot_media.stem}__{fp_a[:8]}_{fp_b[:8]}.json"

        if cache_path.exists():
            eprint(f"[cache] {metric}: using cached pairwise annotations")
            obj = PairwiseAnnotation.model_validate(json.loads(cache_path.read_text(encoding="utf-8")))
        else:
            eprint(f"[gemini] annotating metric={metric} | A={top_media.name} vs B={bot_media.name}")
            top_signals = signals_map.get(metric, [])
            obj = annotate_pairwise(
                client=client,
                model=args.model_annotate,
                metric=metric,
                a_img_path=top_img,
                b_img_path=bot_img,
                a_ctx=a_ctx,
                b_ctx=b_ctx,
                top_signals=top_signals,
            )
            cache_path.write_text(obj.model_dump_json(indent=2), encoding="utf-8")

        # Write annotation JSON and Markdown
        ann_json_path = ann_dir / f"{metric}__A_{top_media.stem}__B_{bot_media.stem}.json"
        ann_md_path = ann_dir / f"{metric}__A_{top_media.stem}__B_{bot_media.stem}.md"
        safe_mkdir(ann_json_path.parent)
        ann_json_path.write_text(obj.model_dump_json(indent=2), encoding="utf-8")

        # Markdown with longer descriptions
        lines = []
        lines.append(f"# Pairwise creative annotations — metric: {metric}\n")
        lines.append(f"## A (higher): {top_media.name}\n")
        lines.append(f"- CTR: {format_metric(top.get('click_through_rate'))}\n")
        lines.append(f"- CVR: {format_metric(top.get('conversion_rate'))}\n")
        lines.append(f"- ATC: {format_metric(top.get('add_to_cart_rate'))}\n")
        lines.append(f"- Media kind: {top_kind}\n\n")

        lines.append(f"## B (lower): {bot_media.name}\n")
        lines.append(f"- CTR: {format_metric(bot.get('click_through_rate'))}\n")
        lines.append(f"- CVR: {format_metric(bot.get('conversion_rate'))}\n")
        lines.append(f"- ATC: {format_metric(bot.get('add_to_cart_rate'))}\n")
        lines.append(f"- Media kind: {bot_kind}\n\n")

        lines.append("## Callouts\n")
        for c in obj.callouts:
            mf = ", ".join(c.metric_focus)
            lines.append(f"- **{c.id}** ({c.target}, impacts: {mf}, {c.direction}) — **{c.label}**: {c.why}\n")
        if obj.notes:
            lines.append(f"\n## Notes\n{obj.notes}\n")
        ann_md_path.write_text("".join(lines), encoding="utf-8")

        # Draw annotated images
        callouts = obj.callouts
        a_out = img_dir / metric / f"A__{top_media.stem}.png"
        b_out = img_dir / metric / f"B__{bot_media.stem}.png"

        # Green for A (higher), Red for B (lower)
        draw_callouts(top_img, callouts, target="A", out_path=a_out, circle_color=(0, 180, 90))
        draw_callouts(bot_img, callouts, target="B", out_path=b_out, circle_color=(220, 60, 60))

        metric_sections.append({
            "metric": metric,
            "A": {
                "file_name": top_media.name,
                "annotated_image": a_out,
                "ctr": format_metric(top.get("click_through_rate")),
                "cvr": format_metric(top.get("conversion_rate")),
                "atc": format_metric(top.get("add_to_cart_rate")),
            },
            "B": {
                "file_name": bot_media.name,
                "annotated_image": b_out,
                "ctr": format_metric(bot.get("click_through_rate")),
                "cvr": format_metric(bot.get("conversion_rate")),
                "atc": format_metric(bot.get("add_to_cart_rate")),
            },
            "callouts": [c.model_dump() for c in callouts],
            "notes": obj.notes,
            "annotation_json": ann_json_path,
            "annotation_md": ann_md_path,
        })

    # Build docx
    out_docx = out_dir / args.docx_name
    if not args.no_docx and metric_sections:
        try:
            build_docx(
                out_docx=out_docx,
                title="Annotated Creative Performance Report",
                dataset_summary=dataset_summary,
                metric_sections=metric_sections,
            )
            eprint(f"[ok] wrote DOCX: {out_docx}")
        except Exception as e:
            eprint(f"[warn] failed to create DOCX: {e}")

    # Print outputs
    print("\nOutputs:")
    print(f"- Annotated images: {img_dir}")
    print(f"- Callout files: {ann_dir}")
    if (not args.no_docx) and out_docx.exists():
        print(f"- Word doc: {out_docx}")
    print("")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())

